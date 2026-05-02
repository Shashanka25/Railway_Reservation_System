import docx
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

doc = docx.Document()

# Title
title = doc.add_heading('Railway Reservation Project Summary', 0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Introduction
doc.add_heading('Overview', level=1)
p1 = doc.add_paragraph(
    "This project is a modern, high-speed railway reservation web application built using Next.js and Tailwind CSS. "
    "It provides a sleek, responsive, and intuitive user interface designed to elevate the ticket booking experience for Indian Railways, "
    "drawing inspiration from premium travel and high-speed rail services like the Vande Bharat Express."
)

# Key Features
doc.add_heading('Key Features', level=1)

features = [
    ("Dynamic Search & Navigation", "Users can search for trains between specific cities and select travel dates. The inputs are dynamically passed as URL parameters to render personalized search results."),
    ("Rich Aesthetics & UI", "The application employs a premium design language. It incorporates a modern color palette, glassmorphism, dynamic gradients, smooth hover effects, and a highly responsive layout."),
    ("Comprehensive Train Listings", "The search results page displays available trains (e.g., RailConnect Express, Rajdhani, Shatabdi) featuring accurate durations, dynamically calculated travel distances, class pricing (EC, CC, 1A, etc.), and real-time-like seat availability."),
    ("Seamless Booking Flow", "The project includes a streamlined user journey spanning from the Home page to Search Results, Seat Selection, Review, and final Ticket confirmation."),
    ("Modern Tech Stack", "Built on Next.js (App Router) using React 19 and Tailwind CSS v4, utilizing both Server and Client Components for optimized performance and interactivity.")
]

for title_text, desc_text in features:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(title_text + ": ")
    run.bold = True
    p.add_run(desc_text)

# Libraries Used
doc.add_heading('Libraries Used', level=1)
libraries = [
    ("Next.js (v16.2)", "The core React framework used for server-side rendering, routing, and overall application architecture."),
    ("React & React DOM (v19.2)", "The foundational UI library for building the interactive user interface."),
    ("Tailwind CSS (v4)", "A utility-first CSS framework used for styling the entire application with premium aesthetics and responsive layouts."),
    ("TypeScript", "Used for static type checking across the codebase to ensure robustness and code quality.")
]
for lib_name, lib_desc in libraries:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(lib_name + ": ")
    run.bold = True
    p.add_run(lib_desc)

# Algorithms Implemented
doc.add_heading('Algorithms Implemented', level=1)
algorithms = [
    ("String Hashing Algorithm (Mock Distance Calculator)", "A custom polynomial rolling hash-style algorithm implemented in the search page. It calculates a consistent integer hash based on the 'From' and 'To' station codes' ASCII character values, transforming the string data into a deterministic pseudo-random distance (ranging from 300km to 1800km) to display on the UI.")
]
for alg_name, alg_desc in algorithms:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(alg_name + ": ")
    run.bold = True
    p.add_run(alg_desc)

# Conclusion
doc.add_heading('Conclusion', level=1)
p2 = doc.add_paragraph(
    "Overall, this project serves as a cutting-edge frontend prototype for a next-generation railway booking platform, "
    "combining aesthetic excellence with functional, dynamic routing to create a seamless user experience."
)

doc.save('Railway_Reservation_Summary_v2.docx')
print("Document saved successfully.")
